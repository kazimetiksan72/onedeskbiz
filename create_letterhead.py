#!/usr/bin/env python3
"""PXL Antetli Kağıt Oluşturucu"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH = "/Users/kazimetiksan/Downloads/PXL (2000 x 1000 piksel) (2000 x 900 piksel) (2000 x 800 piksel).png"
OUTPUT_PATH = str(Path(__file__).resolve().parent / "PXL_Antetli_Kagit.docx")

C_PURPLE = "7C3AED"
C_CYAN   = "00B8D9"
C_PINK   = "EC4899"
C_DARK   = "1F2937"
C_GRAY   = "6B7280"


def rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def clear_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        tc = row.cells[col_idx]._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width_cm * 567)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(height_cm * 567)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def add_color_stripe(container, colors, height_cm=0.22):
    table = container.add_table(rows=1, cols=len(colors), width=Cm(16.0))
    clear_table_borders(table)
    set_row_height(table.rows[0], height_cm)
    widths = [5.33, 5.33, 5.34]  # 16 cm total / 3 cols
    for i, (color, w) in enumerate(zip(colors, widths)):
        cell = table.cell(0, i)
        set_cell_bg(cell, color)
        clear_cell_borders(cell)
        set_col_width(table, i, w)
        p = cell.paragraphs[0]
        p.clear()
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)


def no_space(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


# ─── Document ────────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]

section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Cm(4.8)
section.bottom_margin = Cm(4.2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.header_distance = Cm(0.8)
section.footer_distance = Cm(0.8)

# ─── HEADER ──────────────────────────────────────────────────────────────────
header = section.header

# Remove the default empty paragraph
for p in list(header.paragraphs):
    p._element.getparent().remove(p._element)

# Logo | Company info — 2 column table
ht = header.add_table(rows=1, cols=2, width=Cm(16.0))
clear_table_borders(ht)
set_col_width(ht, 0, 4.5)
set_col_width(ht, 1, 11.5)

# Logo cell
logo_cell = ht.cell(0, 0)
clear_cell_borders(logo_cell)
logo_para = logo_cell.paragraphs[0]
logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
no_space(logo_para)
logo_run = logo_para.add_run()
logo_run.add_picture(LOGO_PATH, width=Cm(4.0))

# Info cell
info_cell = ht.cell(0, 1)
clear_cell_borders(info_cell)
info_para = info_cell.paragraphs[0]
info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
no_space(info_para)

r1 = info_para.add_run("Piksel Teknoloji Çözümleri AŞ")
r1.bold = True
r1.font.size = Pt(11)
r1.font.color.rgb = rgb(C_PURPLE)

info_para.add_run("\n")

r2 = info_para.add_run("pikselmutfak.com")
r2.font.size = Pt(8.5)
r2.font.color.rgb = rgb(C_GRAY)

# Spacer paragraph before stripe
sp = header.add_paragraph()
no_space(sp)
sp_run = sp.add_run(" ")
sp_run.font.size = Pt(3)

# Color stripe
add_color_stripe(header, [C_PURPLE, C_CYAN, C_PINK])

# ─── FOOTER ──────────────────────────────────────────────────────────────────
footer = section.footer

for p in list(footer.paragraphs):
    p._element.getparent().remove(p._element)

# Color stripe
add_color_stripe(footer, [C_PURPLE, C_CYAN, C_PINK])

# Spacer
sp2 = footer.add_paragraph()
no_space(sp2)
sp2.add_run(" ").font.size = Pt(3)

# Contact info
fp = footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(fp)

line1 = (
    "Acıbadem Mah. Çeçen Sk. Akasya AVM No:25 İç Kapı No:150 Üsküdar / İstanbul"
    "   |   T: 0530 179 31 50   |   info@pikselmutfak.com"
)
line2 = "Vergi Dairesi: Üsküdar   |   Vergi No: 7291330999   |   Ticaret Sicil No: 1062882"

r_f1 = fp.add_run(line1 + "\n" + line2)
r_f1.font.size = Pt(7.5)
r_f1.font.color.rgb = rgb(C_GRAY)

# ─── BODY ────────────────────────────────────────────────────────────────────
def body_label(doc, label, line_len=35):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(C_DARK)
    r2 = p.add_run("_" * line_len)
    r2.font.size = Pt(10)
    r2.font.color.rgb = rgb(C_GRAY)
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "60")
    pPr.append(spacing)


body_label(doc, "Tarih:  ")
body_label(doc, "Konu:  ")

spacer = doc.add_paragraph()
spacer._p.get_or_add_pPr().append(
    OxmlElement("w:spacing")
)
spacer._p.pPr.find(qn("w:spacing")).set(qn("w:before"), "120")
spacer._p.pPr.find(qn("w:spacing")).set(qn("w:after"), "0")

content_lines = [
    "Sayın İlgili,",
    "",
    "",
    "",
    "",
    "Saygılarımızla,",
    "",
    "",
    "",
    "Piksel Teknoloji Çözümleri AŞ",
]
for line in content_lines:
    p = doc.add_paragraph(line)
    r = p.runs[0] if p.runs else p.add_run()
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb(C_DARK)
    pPr = p._p.get_or_add_pPr()
    sp_el = OxmlElement("w:spacing")
    sp_el.set(qn("w:before"), "0")
    sp_el.set(qn("w:after"), "80")
    pPr.append(sp_el)

doc.save(OUTPUT_PATH)
print(f"✓ Oluşturuldu: {OUTPUT_PATH}")
